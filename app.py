import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
from io import BytesIO

# =========================================================
# PAGE CONFIG
# =========================================================
st.set_page_config(
    page_title="Sprih Internal Automations",
    page_icon="🚀",
    layout="wide",
)

PRIMARY_GREEN = "1A3A2A"
AMBER = "F59E0B"


# =========================================================
# SESSION STATE
# =========================================================
if "coach_reports" not in st.session_state:
    st.session_state.coach_reports = []

if "gtm_reports" not in st.session_state:
    st.session_state.gtm_reports = []


# =========================================================
# HELPERS
# =========================================================
def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def generate_call_coach_report(data):
    """
    Mimics call coach skill structure
    """
    transcript = data["transcript"].lower()

    # Dynamic scoring
    meddic = {
        "Metrics": 8 if "cost" in transcript else 6,
        "Economic Buyer": 7 if "stakeholder" in transcript else 5,
        "Decision Criteria": 8 if "requirement" in transcript else 6,
        "Decision Process": 7,
        "Identify Pain": 9 if "pain" in transcript else 6,
        "Champion": 5,
    }

    deal_health = {
        "Discovery Depth": "Strong",
        "Pain Clarity": "Adequate",
        "Deal Control": "Developing",
        "Next Steps": "Adequate",
    }

    actions = [
        "TODAY — send quantified follow-up questions",
        "THIS WEEK — map stakeholder decision process",
        "BEFORE NEXT CALL — prepare ROI narrative",
    ]

    return {
        "header": data,
        "classification": (
            f"This was classified as a {data['call_type']} call "
            f"originating from {data['origination']} based on the "
            f"transcript signals and current deal stage "
            f"({data['deal_stage']})."
        ),
        "meddic": meddic,
        "deal_health": deal_health,
        "actions": actions,
        "coaching_moments": [
            {
                "what": "You uncovered the problem clearly.",
                "why": "This creates strong discovery momentum.",
                "better": (
                    "Can you help quantify the business impact?"
                ),
            },
            {
                "what": "You did not probe decision authority deeply.",
                "why": (
                    "This is an opportunity to improve deal control."
                ),
                "better": (
                    "Who besides you would need to sign off?"
                ),
            },
        ],
        "themes": {
            "signals": [
                "budget sensitivity",
                "timeline pressure",
            ],
            "requirements": [
                "faster reporting",
                "cross-functional visibility",
            ],
            "patterns": [
                "repeated manual process pain",
            ],
        },
        "overall_score": round(sum(meddic.values()) / len(meddic), 1),
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
    }


def create_docx_report(report):
    doc = Document()

    # =====================================================
    # HEADER
    # =====================================================
    header = doc.add_table(rows=1, cols=2)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER

    left = header.cell(0, 0)
    right = header.cell(0, 1)

    shade_cell(left, PRIMARY_GREEN)
    shade_cell(right, AMBER)

    header_data = report["header"]

    run = left.paragraphs[0].add_run(
        f"CALL COACH REPORT\n"
        f"{header_data['rep_name']}\n"
        f"{header_data['company']}\n"
        f"{header_data['call_date']}\n"
        f"{header_data['call_type']} | "
        f"{header_data['origination']}"
    )
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.size = Pt(13)
    run.bold = True

    score_run = right.paragraphs[0].add_run(
        f"Overall\n{report['overall_score']}/10"
    )
    score_run.font.size = Pt(16)
    score_run.bold = True

    doc.add_paragraph()

    # =====================================================
    # CONTEXT SUMMARY
    # =====================================================
    doc.add_heading("1. CONTEXT SUMMARY", level=1)

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"

    rows = [
        ("Call Type", header_data["call_type"]),
        ("Origination", header_data["origination"]),
        (
            "Discovery Previously Done",
            header_data["discovery_done"],
        ),
        ("Deal Stage", header_data["deal_stage"]),
        ("Previous Notes", header_data["previous_notes"]),
    ]

    for i, (k, v) in enumerate(rows):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v

    doc.add_paragraph(
        f"Calibration Note: {header_data['calibration_note']}"
    )

    # =====================================================
    # CALL CLASSIFICATION
    # =====================================================
    doc.add_heading("2. CALL CLASSIFICATION", level=1)
    doc.add_paragraph(report["classification"])

    # =====================================================
    # FRAMEWORK SCORES
    # =====================================================
    doc.add_heading("3. FRAMEWORK SCORES", level=1)

    score_table = doc.add_table(
        rows=len(report["meddic"]) + 1,
        cols=3
    )
    score_table.style = "Table Grid"

    score_table.cell(0, 0).text = "Dimension"
    score_table.cell(0, 1).text = "Score"
    score_table.cell(0, 2).text = "Path Forward"

    for i, (k, v) in enumerate(
        report["meddic"].items(), start=1
    ):
        score_table.cell(i, 0).text = k
        score_table.cell(i, 1).text = f"{v}/10"

        if v <= 6:
            score_table.cell(i, 2).text = (
                "Probe deeper in next interaction"
            )
        else:
            score_table.cell(i, 2).text = "Maintain strength"

    # =====================================================
    # DEAL HEALTH
    # =====================================================
    doc.add_heading("4. DEAL HEALTH", level=1)

    for k, v in report["deal_health"].items():
        doc.add_paragraph(f"{k}: {v}")

    # =====================================================
    # COACHING MOMENTS
    # =====================================================
    doc.add_heading("5. COACHING MOMENTS", level=1)

    for idx, moment in enumerate(
        report["coaching_moments"], start=1
    ):
        doc.add_paragraph(
            f"Moment {idx} — What happened: {moment['what']}"
        )
        doc.add_paragraph(
            f"Why it matters: {moment['why']}"
        )
        doc.add_paragraph(
            f"Better version: {moment['better']}"
        )

    # =====================================================
    # THEMES
    # =====================================================
    doc.add_heading("6. THEMES & INTELLIGENCE", level=1)

    doc.add_paragraph("Signals & Objections")
    for x in report["themes"]["signals"]:
        doc.add_paragraph(f"• {x}")

    doc.add_paragraph("Requirements")
    for x in report["themes"]["requirements"]:
        doc.add_paragraph(f"• {x}")

    doc.add_paragraph("Patterns")
    for x in report["themes"]["patterns"]:
        doc.add_paragraph(f"• {x}")

    # =====================================================
    # ACTIONS
    # =====================================================
    doc.add_heading("7. TOP 3 IMMEDIATE ACTIONS", level=1)

    for action in report["actions"]:
        doc.add_paragraph(action)

    # FOOTER
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        "Sprih Call Coach (automated)"
    )
    footer_run.italic = True
    footer_run.font.size = Pt(10)

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)

    return stream


# =========================================================
# MAIN HEADER
# =========================================================
st.title("🚀 Sprih Internal Automations")
st.caption(
    "Call Coach + GTM Outreach internal workflow platform"
)

# =========================================================
# SIDEBAR (HTML-like structure)
# =========================================================
st.sidebar.title("Automations")

module = st.sidebar.radio(
    "Select",
    [
        "📞 Call Coach",
        "🎯 GTM Outreach",
        "📊 Dashboard",
        "📚 History",
    ],
)

# =========================================================
# DASHBOARD
# =========================================================
if module == "📊 Dashboard":
    st.header("📊 Dashboard")

    col1, col2 = st.columns(2)

    col1.metric(
        "Call Coach Reports",
        len(st.session_state.coach_reports),
    )

    col2.metric(
        "GTM Runs",
        len(st.session_state.gtm_reports),
    )

# =========================================================
# CALL COACH
# =========================================================
elif module == "📞 Call Coach":
    st.header("📞 Call Coach")

    col1, col2 = st.columns([2, 1])

    with col1:
        rep_name = st.text_input("Sales Rep Name")
        company = st.text_input("Prospect / Company")
        call_date = st.date_input("Call Date")

        call_type = st.selectbox(
            "Call Type",
            [
                "Discovery",
                "Demo",
                "Follow-up",
                "Commercial",
            ],
        )

        origination = st.selectbox(
            "Origination",
            [
                "Inbound",
                "Outbound",
                "Referral",
            ],
        )

        deal_stage = st.selectbox(
            "Deal Stage",
            [
                "Discovery",
                "Qualification",
                "Proposal",
                "Negotiation",
            ],
        )

        discovery_done = st.selectbox(
            "Discovery Previously Done",
            ["Yes", "No", "Partial"],
        )

        previous_notes = st.text_area(
            "Previous Call Notes"
        )

        calibration_note = st.text_area(
            "Calibration Note"
        )

        transcript = st.text_area(
            "Transcript",
            height=250,
        )

        generate = st.button("🎯 Generate Report")

    with col2:
        st.info(
            """
**Frameworks**
- MEDDIC
- Deal Health
- Coaching Moments
- Themes
- Immediate Actions
"""
        )

    if generate:
        data = {
            "rep_name": rep_name,
            "company": company,
            "call_date": str(call_date),
            "call_type": call_type,
            "origination": origination,
            "deal_stage": deal_stage,
            "discovery_done": discovery_done,
            "previous_notes": previous_notes,
            "calibration_note": calibration_note,
            "transcript": transcript,
        }

        report = generate_call_coach_report(data)

        st.session_state.coach_reports.append(report)

        st.success("Report generated")

        st.subheader("Report Preview")

        st.metric(
            "Overall Score",
            f"{report['overall_score']}/10",
        )

        for k, v in report["meddic"].items():
            st.write(f"**{k}**: {v}/10")

        docx = create_docx_report(report)

        st.download_button(
            "📥 Download DOCX Report",
            docx,
            file_name=f"{rep_name}_call_coach_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

# =========================================================
# GTM OUTREACH
# =========================================================
elif module == "🎯 GTM Outreach":
    st.header("🎯 GTM Outreach")

    company_name = st.text_input("Target Company")
    persona = st.text_input("Persona")
    hook = st.text_area("Hook / Personalisation")

    if st.button("Generate Outreach Brief"):
        st.success("Outreach brief created")

        st.write("### Sample Email")
        st.write(
            f"Hi team at {company_name},\n\n"
            f"I noticed {hook}.\n"
            f"I’d love to connect regarding ESG workflows."
        )

# =========================================================
# HISTORY
# =========================================================
elif module == "📚 History":
    st.header("📚 History")

    for report in reversed(
        st.session_state.coach_reports
    ):
        with st.expander(
            report["header"]["rep_name"]
        ):
            st.write(report)
